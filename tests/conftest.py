import pytest
import pandas as pd
import openpyxl
from pathlib import Path


@pytest.fixture
def simple_xlsx(tmp_path) -> Path:
    """Single-sheet xlsx with clean data."""
    path = tmp_path / "simple.xlsx"
    df = pd.DataFrame({"Name": ["Alice", "Bob"], "Age": [30, 25], "Score": [95.5, 87.0]})
    df.to_excel(path, index=False)
    return path


@pytest.fixture
def merged_cells_xlsx(tmp_path) -> Path:
    """Xlsx with horizontally and vertically merged cells."""
    path = tmp_path / "merged.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Report"
    ws["A1"] = "Region"
    ws["B1"] = "Q1"
    ws["C1"] = "Q1"
    ws.merge_cells("B1:C1")
    ws["A2"] = "North"
    ws["B2"] = 100
    ws["C2"] = 200
    ws.merge_cells("A3:A4")
    ws["A3"] = "South"
    ws["B3"] = 150
    ws["C3"] = 175
    ws["B4"] = 120
    ws["C4"] = 130
    wb.save(path)
    return path


@pytest.fixture
def multi_sheet_xlsx(tmp_path) -> Path:
    """Xlsx with two sheets."""
    path = tmp_path / "multi.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"A": [1, 2]}).to_excel(writer, sheet_name="Sheet1", index=False)
        pd.DataFrame({"X": [9, 8]}).to_excel(writer, sheet_name="Sheet2", index=False)
    return path


@pytest.fixture
def csv_comma(tmp_path) -> Path:
    path = tmp_path / "data.csv"
    path.write_text("Name,Age,City\nAlice,30,Hanoi\nBob,25,HCMC\n", encoding="utf-8")
    return path


@pytest.fixture
def csv_semicolon(tmp_path) -> Path:
    path = tmp_path / "data_semi.csv"
    path.write_text("Name;Age;City\nAlice;30;Hanoi\nBob;25;HCMC\n", encoding="utf-8")
    return path


@pytest.fixture
def csv_utf16(tmp_path) -> Path:
    path = tmp_path / "data_utf16.csv"
    path.write_bytes("Name,Value\nAlpha,1\nBeta,2\n".encode("utf-16"))
    return path


@pytest.fixture
def simple_docx(tmp_path) -> Path:
    """Simple docx with heading and paragraph."""
    from docx import Document
    path = tmp_path / "letter.docx"
    doc = Document()
    doc.add_heading("Hello World", level=1)
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.save(str(path))
    return path


@pytest.fixture
def digital_pdf(tmp_path) -> Path:
    """PDF with embedded text layer (digital, not scanned)."""
    from fpdf import FPDF
    path = tmp_path / "digital.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for i in range(20):
        pdf.cell(0, 10, f"Line {i}: This is digital text content for testing.", ln=True)
    pdf.output(str(path))
    return path
